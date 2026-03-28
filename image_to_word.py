from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class ImageToWordConverter:
    def __init__(self, tesseract_path=None):
        """Initialize the converter with optional Tesseract path."""
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    def extract_text_from_image(self, image_path):
        """Extract text from image using OCR."""
        try:
            image = Image.open(image_path)
            # Use Tesseract to extract text
            text = pytesseract.image_to_string(image, lang='eng')
            return text
        except Exception as e:
            raise Exception(f"Error extracting text: {e}")
    
    def create_word_document(self, text, output_path, include_formatting=True):
        """Create a Word document from extracted text."""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extracted Text from Image', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add extracted text
        if include_formatting:
            # Split by paragraphs and add them
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.style.font.size = Pt(11)
        else:
            # Add as single block
            doc.add_paragraph(text)
        
        # Save document
        doc.save(output_path)
        print(f"Word document saved: {output_path}")
    
    def convert_image_to_word(self, image_path, output_path=None):
        """Convert image to Word document."""
        if output_path is None:
            base_name = os.path.splitext(image_path)[0]
            output_path = f"{base_name}.docx"
        
        # Extract text
        print("Extracting text from image...")
        text = self.extract_text_from_image(image_path)
        
        # Create Word document
        print("Creating Word document...")
        self.create_word_document(text, output_path)
        
        return output_path

if __name__ == "__main__":
    # Example usage
    converter = ImageToWordConverter()
    
    # Convert image to Word
    image_file = "dictionary.png"  # Replace with your image path
    output_file = "dictionary_output.docx"
    
    if os.path.exists(image_file):
        converter.convert_image_to_word(image_file, output_file)
    else:
        print(f"Image file not found: {image_file}")
        print("Please provide the path to your image file.")
